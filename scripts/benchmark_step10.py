#!/usr/bin/env python3
from __future__ import annotations

import argparse
import json
import os
import shutil
import subprocess
import sys
import time
from pathlib import Path
from tempfile import TemporaryDirectory

from docx import Document
from fastapi.testclient import TestClient
from PIL import Image
from sqlalchemy import select


def main() -> int:
    parser = argparse.ArgumentParser(description="Step 10 concurrency benchmark")
    parser.add_argument("--tasks", type=int, default=20, help="number of tasks in one batch")
    parser.add_argument("--workers", type=int, default=6, help="number of rq burst workers")
    parser.add_argument("--timeout", type=int, default=120, help="timeout seconds")
    parser.add_argument("--ocr-provider", type=str, default="mock", help="mock|paddleocr")
    parser.add_argument("--llm-provider", type=str, default="mock", help="mock|openrouter")
    args = parser.parse_args()

    repo_root = Path(__file__).resolve().parents[1]
    backend_dir = repo_root / "backend"

    os.environ["OCR_PROVIDER"] = args.ocr_provider
    os.environ["LLM_PROVIDER"] = args.llm_provider

    sys.path.insert(0, str(backend_dir))
    from app.db.session import SessionLocal
    from app.main import app
    from app.models.entities import Task, TaskStatus
    from app.queue.rq_app import get_queue

    client = TestClient(app)
    queue = get_queue()
    queue.empty()

    with TemporaryDirectory() as td:
        td_path = Path(td)
        book_file = td_path / "bench.docx"
        doc = Document()
        doc.add_heading("并发压测书稿", level=1)
        doc.add_paragraph("执行 复盘 方法 结构 输出。" * 200)
        doc.save(str(book_file))

        with book_file.open("rb") as f:
            rb = client.post(
                "/books/upload",
                files={"file": ("bench.docx", f, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
                data={"title": "并发压测书稿"},
            )
        if rb.status_code != 201:
            print("book upload failed:", rb.status_code, rb.text)
            return 1
        book_id = rb.json()["id"]

        # seed random tag pool
        for i in range(1, 21):
            client.post("/tags", json={"tag_text": f"标签{i}", "enabled": True})

        # build one request with N folders
        bindings = []
        files = []
        for i in range(1, args.tasks + 1):
            folder = f"note_{i:02d}"
            bindings.append({"folder_name": folder, "book_id": book_id})
            img_path = td_path / f"{folder}.jpg"
            Image.new("RGB", (120, 40), color="white").save(img_path)
            files.append(("files", (f"{folder}/图1.jpg", img_path.read_bytes(), "image/jpeg")))

        t0 = time.perf_counter()
        rt = client.post(
            "/tasks",
            data={
                "bindings": json.dumps(bindings, ensure_ascii=False),
                "batch_name": "step10-benchmark",
                "auto_enqueue": "true",
            },
            files=files,
        )
        if rt.status_code != 201:
            print("task creation failed:", rt.status_code, rt.text)
            return 1
        task_ids = rt.json()["task_ids"]

        # run workers in parallel burst mode
        env = os.environ.copy()
        env["OCR_PROVIDER"] = args.ocr_provider
        env["LLM_PROVIDER"] = args.llm_provider

        procs: list[subprocess.Popen] = []
        rq_bin = shutil.which("rq")
        if not rq_bin:
            print("rq command not found in current environment.")
            return 1
        for _ in range(args.workers):
            procs.append(
                subprocess.Popen(
                    [
                        rq_bin,
                        "worker",
                        "xhsocr_tasks",
                        "--url",
                        "redis://127.0.0.1:6379/0",
                        "--burst",
                        "--worker-class",
                        "rq.worker.SimpleWorker",
                    ],
                    cwd=str(backend_dir),
                    env=env,
                    stdout=subprocess.DEVNULL,
                    stderr=subprocess.DEVNULL,
                )
            )

        deadline = time.time() + args.timeout
        final_status: dict[int, str] = {}
        while time.time() < deadline:
            with SessionLocal() as db:
                rows = db.execute(select(Task.id, Task.status).where(Task.id.in_(task_ids))).all()
            final_status = {int(r[0]): r[1].value for r in rows}
            done = all(v in (TaskStatus.success.value, TaskStatus.failed.value) for v in final_status.values())
            if done and len(final_status) == len(task_ids):
                break
            time.sleep(0.5)

        for p in procs:
            try:
                p.wait(timeout=5)
            except subprocess.TimeoutExpired:
                p.kill()

        elapsed = time.perf_counter() - t0
        success_count = sum(1 for s in final_status.values() if s == TaskStatus.success.value)
        failed_count = sum(1 for s in final_status.values() if s == TaskStatus.failed.value)
        throughput = (success_count + failed_count) / elapsed if elapsed > 0 else 0

        report = {
            "tasks": args.tasks,
            "workers": args.workers,
            "ocr_provider": args.ocr_provider,
            "llm_provider": args.llm_provider,
            "elapsed_seconds": round(elapsed, 3),
            "success_count": success_count,
            "failed_count": failed_count,
            "throughput_tasks_per_sec": round(throughput, 3),
            "status_map": final_status,
        }
        print(json.dumps(report, ensure_ascii=False, indent=2))
        return 0 if success_count + failed_count == args.tasks and failed_count == 0 else 2


if __name__ == "__main__":
    raise SystemExit(main())
